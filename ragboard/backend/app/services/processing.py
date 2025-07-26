"""
File processing service for content extraction.
"""

from typing import Dict, Any, Optional, Tuple, List
from uuid import UUID
from pathlib import Path
import asyncio
import aiofiles
import logging
from datetime import datetime
import mimetypes
import hashlib

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import whisper
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False

# Web scraping imports
try:
    from bs4 import BeautifulSoup
    import httpx
    WEB_AVAILABLE = True
except ImportError:
    WEB_AVAILABLE = False

try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_AVAILABLE = True
except ImportError:
    YOUTUBE_AVAILABLE = False

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update

from app.core.config import settings
from app.models.resource import Resource, ResourceType, ProcessingStatus
from app.db.base import async_session_maker
from app.services.rag_pipeline import RAGPipeline
from app.db.base import get_async_session

logger = logging.getLogger(__name__)


class ProcessingService:
    """Service for processing various file types and extracting content."""
    
    def __init__(self):
        self.rag_pipeline = RAGPipeline()
        self.whisper_model = None
        
        # Initialize Whisper if available
        if WHISPER_AVAILABLE and hasattr(settings, 'whisper_model'):
            try:
                self.whisper_model = whisper.load_model(settings.whisper_model)
                logger.info(f"Whisper model {settings.whisper_model} loaded")
            except Exception as e:
                logger.error(f"Failed to load Whisper model: {e}")
    
    async def queue_resource(self, resource_id: UUID, force: bool = False):
        """Queue a resource for processing."""
        # In a production environment, this would use a task queue like Celery
        # For now, we'll process directly
        await self.process_resource(resource_id, force)
    
    async def process_resource(self, resource_id: UUID, force: bool = False):
        """
        Process a resource: extract text and generate embeddings.
        
        Args:
            resource_id: Resource ID to process
            force: Force reprocessing even if already processed
        """
        async with async_session_maker() as db:
            # Get resource
            result = await db.execute(
                select(Resource).where(Resource.id == resource_id)
            )
            resource = result.scalar_one_or_none()
            
            if not resource:
                logger.error(f"Resource {resource_id} not found")
                return
            
            # Skip if already processed unless forced
            if resource.processing_status == ProcessingStatus.COMPLETED and not force:
                logger.info(f"Resource {resource_id} already processed")
                return
            
            # Update status to processing
            resource.processing_status = ProcessingStatus.PROCESSING
            await db.commit()
            
            try:
                # Extract text based on resource type
                extracted_text, metadata = await self.extract_text(resource)
                
                if not extracted_text:
                    raise ValueError("No text extracted from resource")
                
                # Update resource with extracted content
                resource.extracted_text = extracted_text
                resource.extracted_metadata = metadata
                
                # Generate embeddings and store in vector DB
                chunk_count = await self.rag_pipeline.process_and_store_resource(
                    resource_id=resource.id,
                    text=extracted_text,
                    chunk_size=1000,
                    chunk_overlap=200
                )
                
                # Update resource status
                resource.processing_status = ProcessingStatus.COMPLETED
                resource.processed_at = datetime.utcnow()
                resource.chunk_count = chunk_count
                
                await db.commit()
                logger.info(f"Successfully processed resource {resource_id} with {chunk_count} chunks")
                
            except Exception as e:
                logger.error(f"Error processing resource {resource_id}: {e}")
                resource.processing_status = ProcessingStatus.FAILED
                resource.processing_error = str(e)
                await db.commit()
    
    async def extract_text(self, resource: Resource) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a resource based on its type.
        
        Args:
            resource: Resource object
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        
        if resource.resource_type == ResourceType.PDF:
            return await self._extract_pdf(resource.file_path)
        elif resource.resource_type == ResourceType.DOCUMENT:
            return await self._extract_document(resource.file_path)
        elif resource.resource_type == ResourceType.IMAGE:
            return await self._extract_image(resource.file_path)
        elif resource.resource_type == ResourceType.AUDIO:
            return await self._extract_audio(resource.file_path)
        elif resource.resource_type == ResourceType.VIDEO:
            return await self._extract_video(resource.file_path, resource.source_url)
        elif resource.resource_type == ResourceType.WEBPAGE:
            return await self._extract_webpage(resource.source_url)
        elif resource.resource_type == ResourceType.TEXT:
            return await self._extract_text_file(resource.file_path or resource.extracted_text)
        else:
            raise ValueError(f"Unsupported resource type: {resource.resource_type}")
    
    async def _extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed")
        
        text_parts = []
        metadata = {"pages": 0, "extracted_pages": []}
        
        async with aiofiles.open(file_path, 'rb') as file:
            content = await file.read()
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            
            metadata["pages"] = len(pdf_reader.pages)
            
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"[Page {i+1}]\n{page_text}")
                        metadata["extracted_pages"].append(i+1)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {i+1}: {e}")
        
        return "\n\n".join(text_parts), metadata
    
    async def _extract_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from document files (DOCX, TXT, etc.)."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in ['.txt', '.md', '.csv']:
            # Plain text files
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            return text, {"file_type": file_ext}
        
        elif file_ext in ['.docx'] and DOCX_AVAILABLE:
            # Word documents
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            text = "\n\n".join(paragraphs)
            if tables_text:
                text += "\n\nTables:\n" + "\n\n".join(tables_text)
            
            return text, {"file_type": "docx", "paragraphs": len(paragraphs), "tables": len(doc.tables)}
        
        else:
            # Fallback: try to read as text
            try:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                    text = await file.read()
                return text, {"file_type": file_ext, "fallback": True}
            except Exception as e:
                raise ValueError(f"Cannot extract text from {file_ext} file: {e}")
    
    async def _extract_image(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image using OCR."""
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract and Pillow not installed")
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang=settings.ocr_languages)
            
            metadata = {
                "image_size": image.size,
                "image_mode": image.mode,
                "ocr_language": settings.ocr_languages
            }
            
            return text.strip(), metadata
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            return "", {"error": str(e)}
    
    async def _extract_audio(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from audio using Whisper."""
        if not WHISPER_AVAILABLE or not self.whisper_model:
            raise ImportError("Whisper not available")
        
        try:
            result = self.whisper_model.transcribe(file_path)
            
            metadata = {
                "duration": result.get("duration"),
                "language": result.get("language"),
                "segments": len(result.get("segments", []))
            }
            
            return result["text"], metadata
        except Exception as e:
            logger.error(f"Audio transcription failed for {file_path}: {e}")
            raise
    
    async def _extract_video(self, file_path: Optional[str], source_url: Optional[str]) -> Tuple[str, Dict[str, Any]]:
        """Extract text from video (YouTube transcripts or audio extraction)."""
        # Check if it's a YouTube video
        if source_url and "youtube.com" in source_url or "youtu.be" in source_url:
            return await self._extract_youtube_transcript(source_url)
        
        # For other videos, extract audio and transcribe
        if file_path and WHISPER_AVAILABLE:
            # In production, you would extract audio track first
            # For now, we'll try to transcribe directly
            return await self._extract_audio(file_path)
        
        raise ValueError("Cannot extract text from video")
    
    async def _extract_youtube_transcript(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Extract transcript from YouTube video."""
        if not YOUTUBE_AVAILABLE:
            raise ImportError("youtube-transcript-api not installed")
        
        try:
            # Extract video ID
            video_id = None
            if "v=" in url:
                video_id = url.split("v=")[1].split("&")[0]
            elif "youtu.be/" in url:
                video_id = url.split("youtu.be/")[1].split("?")[0]
            
            if not video_id:
                raise ValueError("Could not extract video ID from URL")
            
            # Get transcript
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            
            # Format transcript
            text_parts = []
            for entry in transcript_list:
                text_parts.append(entry['text'])
            
            metadata = {
                "video_id": video_id,
                "duration": transcript_list[-1]['start'] + transcript_list[-1]['duration'] if transcript_list else 0,
                "segments": len(transcript_list)
            }
            
            return " ".join(text_parts), metadata
            
        except Exception as e:
            logger.error(f"Failed to get YouTube transcript: {e}")
            raise
    
    async def _extract_webpage(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from webpage."""
        if not WEB_AVAILABLE:
            raise ImportError("beautifulsoup4 and httpx not installed")
        
        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()
                
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract metadata
                metadata = {
                    "url": str(response.url),
                    "title": soup.title.string if soup.title else None,
                    "description": None,
                    "keywords": None
                }
                
                # Get meta tags
                for meta in soup.find_all('meta'):
                    if meta.get('name') == 'description':
                        metadata['description'] = meta.get('content')
                    elif meta.get('name') == 'keywords':
                        metadata['keywords'] = meta.get('content')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text
                text = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text, metadata
                
        except Exception as e:
            logger.error(f"Failed to extract webpage content: {e}")
            raise
    
    async def _extract_text_file(self, content: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from text content."""
        if isinstance(content, str):
            return content, {"source": "direct_text"}
        
        # If it's a file path
        async with aiofiles.open(content, 'r', encoding='utf-8') as file:
            text = await file.read()
        
        return text, {"source": "text_file"}


# Add missing import
import io